from __future__ import annotations

import calendar
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path


RESTRICTED_PATTERNS = (
    "保密",
    "勿转",
    "不外传",
    "仅限内部",
    "内部交流",
    "传播限制",
    "未经许可",
    "禁止转发",
)


@dataclass
class ParsedChunk:
    ordinal: int
    text: str
    heading: str | None = None
    page: int | None = None
    line_start: int | None = None
    line_end: int | None = None


@dataclass
class ParsedDocument:
    title: str
    chunks: list[ParsedChunk]
    extraction_status: str
    cloud_allowed: bool
    restriction_reason: str | None


def restriction_reason(text: str) -> str | None:
    matches = [marker for marker in RESTRICTED_PATTERNS if marker in text]
    return "命中受限标记：" + "、".join(matches) if matches else None


def clean_unicode(text: str) -> tuple[str, bool]:
    """Replace malformed surrogate code points before SQLite/JSON serialization."""
    cleaned = text.encode("utf-8", errors="replace").decode("utf-8")
    return cleaned, cleaned != text


def _split_long_text(text: str, max_chars: int = 2600, overlap: int = 180) -> list[str]:
    text = text.strip()
    if not text:
        return []
    if len(text) <= max_chars:
        return [text]
    parts: list[str] = []
    cursor = 0
    while cursor < len(text):
        end = min(len(text), cursor + max_chars)
        if end < len(text):
            boundary = max(text.rfind("\n\n", cursor, end), text.rfind("。", cursor, end))
            if boundary > cursor + max_chars // 2:
                end = boundary + 1
        parts.append(text[cursor:end].strip())
        if end >= len(text):
            break
        cursor = max(cursor + 1, end - overlap)
    return [part for part in parts if part]


def parse_markdown(path: Path) -> ParsedDocument:
    text, had_invalid_unicode = clean_unicode(path.read_text(encoding="utf-8-sig", errors="replace"))
    lines = text.splitlines()
    title = path.stem
    current_heading: str | None = None
    block_start = 1
    block_lines: list[str] = []
    chunks: list[ParsedChunk] = []

    def flush(end_line: int) -> None:
        nonlocal block_lines, block_start
        body = "\n".join(block_lines).strip()
        if body:
            for part in _split_long_text(body):
                chunks.append(
                    ParsedChunk(
                        ordinal=len(chunks),
                        text=part,
                        heading=current_heading,
                        line_start=block_start,
                        line_end=end_line,
                    )
                )
        block_lines = []

    in_frontmatter = bool(lines and lines[0].strip() == "---")
    for index, line in enumerate(lines, start=1):
        if in_frontmatter:
            if index > 1 and line.strip() == "---":
                in_frontmatter = False
                block_start = index + 1
            continue
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            flush(index - 1)
            current_heading = match.group(2).strip()
            if match.group(1) == "#" and title == path.stem:
                title = re.sub(r"[*_`]+", "", current_heading)
            block_start = index
            block_lines = [line]
        else:
            if not block_lines:
                block_start = index
            block_lines.append(line)
    flush(len(lines))

    reason = restriction_reason(text)
    status = "suspect" if "\ufffd" in text or had_invalid_unicode else "ok"
    if status == "suspect" and not reason:
        reason = "文本包含无法解码字符，仅供本地检索"
    return ParsedDocument(title, chunks, status, reason is None and status == "ok", reason)


def _pdf_text_is_suspect(text: str, pages: int) -> bool:
    compact = re.sub(r"\s+", "", text)
    if len(compact) < max(120, pages * 45):
        return True
    replacement_ratio = compact.count("\ufffd") / max(len(compact), 1)
    meaningful = len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", compact)) / max(len(compact), 1)
    return replacement_ratio > 0.005 or meaningful < 0.42


def parse_pdf(path: Path) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("缺少 pypdf，无法提取 PDF") from exc

    reader = PdfReader(str(path))
    page_texts: list[str] = []
    had_invalid_unicode = False
    for page in reader.pages:
        try:
            page_text, page_invalid = clean_unicode((page.extract_text() or "").strip())
            had_invalid_unicode = had_invalid_unicode or page_invalid
            page_texts.append(page_text)
        except Exception:
            page_texts.append("")
    full_text = "\n".join(page_texts)
    if not full_text.strip():
        return ParsedDocument(path.stem, [], "unsupported", False, "PDF 无可提取文本")

    suspect = had_invalid_unicode or _pdf_text_is_suspect(full_text, len(page_texts))
    reason = restriction_reason(full_text)
    chunks: list[ParsedChunk] = []
    for page_no, page_text in enumerate(page_texts, start=1):
        for part in _split_long_text(page_text):
            chunks.append(
                ParsedChunk(
                    ordinal=len(chunks),
                    text=part,
                    heading=f"第 {page_no} 页",
                    page=page_no,
                )
            )
    status = "suspect" if suspect else "ok"
    if suspect and not reason:
        reason = "PDF 文本提取质量存疑，仅供本地检索"
    return ParsedDocument(path.stem, chunks, status, reason is None and not suspect, reason)


def parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，无法提取 DOCX") from exc

    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    text = "\n\n".join(paragraphs)
    if not text:
        return ParsedDocument(path.stem, [], "unsupported", False, "DOCX 无可提取文本")
    reason = restriction_reason(text)
    chunks = [ParsedChunk(index, part) for index, part in enumerate(_split_long_text(text))]
    return ParsedDocument(path.stem, chunks, "ok", reason is None, reason)


def parse_document(path: Path) -> ParsedDocument:
    extension = path.suffix.lower()
    if extension == ".md":
        return parse_markdown(path)
    if extension == ".pdf":
        return parse_pdf(path)
    if extension == ".docx":
        return parse_docx(path)
    if extension in {".png", ".jpg", ".jpeg"}:
        return ParsedDocument(path.stem, [], "unsupported", False, "图片未启用 OCR")
    return ParsedDocument(path.stem, [], "unsupported", False, "不支持的文件类型")


def split_markdown_row(line: str) -> list[str]:
    if not line.lstrip().startswith("|"):
        return []
    placeholder = "\x00PIPE\x00"
    protected = line.strip().strip("|").replace("\\|", placeholder)
    return [cell.strip().replace(placeholder, "|") for cell in protected.split("|")]


def parse_table(path: Path, min_columns: int) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in path.read_text(encoding="utf-8-sig", errors="replace").splitlines():
        cells = split_markdown_row(line)
        if len(cells) < min_columns:
            continue
        if all(re.fullmatch(r":?-{2,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    return rows[1:] if rows else []


def normalize_deadline(raw: str) -> tuple[str | None, str]:
    text = re.sub(r"[*_`]", "", raw)
    candidates: list[tuple[date, str]] = []

    for year, month, day in re.findall(r"(20\d{2})[-/.年](\d{1,2})[-/.月](\d{1,2})", text):
        try:
            candidates.append((date(int(year), int(month), int(day)), "day"))
        except ValueError:
            pass

    for year, quarter in re.findall(r"(20\d{2})\s*[-年]?\s*Q([1-4])", text, re.I):
        month = int(quarter) * 3
        candidates.append((date(int(year), month, calendar.monthrange(int(year), month)[1]), "quarter"))

    for year, half in re.findall(r"(20\d{2})\s*年?\s*(上|下)半年", text):
        month = 6 if half == "上" else 12
        candidates.append((date(int(year), month, calendar.monthrange(int(year), month)[1]), "half-year"))

    for year, month in re.findall(r"(20\d{2})[-/.年](\d{1,2})\s*月", text):
        month_i = int(month)
        if 1 <= month_i <= 12:
            candidates.append((date(int(year), month_i, calendar.monthrange(int(year), month_i)[1]), "month"))

    for year in re.findall(r"(20\d{2})\s*(?:年内|年底|年末|底)", text):
        candidates.append((date(int(year), 12, 31), "year"))

    if not candidates:
        years = re.findall(r"20\d{2}", text)
        if len(years) == 1:
            candidates.append((date(int(years[0]), 12, 31), "year-approx"))
    if not candidates:
        return None, "unknown"
    deadline, precision = max(candidates, key=lambda item: item[0])
    return deadline.isoformat(), precision


def prediction_status(result: str) -> str:
    if "◐" in result or "部分命中" in result:
        return "partial"
    if "❌" in result or "落空" in result:
        return "miss"
    if "✅" in result or ("命中" in result and "待验证" not in result):
        return "hit"
    if "待验证" in result:
        return "pending"
    return "unknown"
